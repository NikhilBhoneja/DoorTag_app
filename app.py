import os
import io
import json
import base64
import tempfile
import requests
import pdfplumber

from flask import Flask, request, send_file, jsonify, render_template
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB max upload

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# ─────────────────────────────────────────────
#  PDF EXTRACTION
# ─────────────────────────────────────────────

def extract_pdf(pdf_path):
    """Extract text from pages 1-2, render pages 3+ as base64 images."""
    text_parts = []
    images_b64 = []

    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        # Pages 1-2: text
        for i in range(min(2, total)):
            t = pdf.pages[i].extract_text() or ""
            text_parts.append(f"[PAGE {i+1}]\n{t}")

    # Pages 3+: render as images for vision
    if total > 2:
        imgs = convert_from_path(pdf_path, dpi=200, first_page=3, last_page=total)
        for img in imgs:
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            images_b64.append(base64.b64encode(buf.getvalue()).decode())

    return "\n\n".join(text_parts), images_b64, total


# ─────────────────────────────────────────────
#  CLAUDE EXTRACTION
# ─────────────────────────────────────────────

def extract_with_claude(order_text, eng_images, api_key):
    """Send order text + engineering sheet images to Claude, get structured JSON."""

    SYSTEM = (
        "You are a door manufacturing data extraction agent for Acme & Dorf Door Corp. "
        "You read door engineering listing sheets (which may be handwritten or typed forms). "
        "Return ONLY a valid JSON object — no markdown fences, no explanation, no extra text."
    )

    content = []

    # If we have engineering sheet images, use vision
    if eng_images:
        content.append({
            "type": "text",
            "text": (
                "Below are engineering sheet page image(s) from a door order PDF "
                "(page 3 onward — these are the handwritten/typed shop forms).\n\n"
                "Also here is the text extracted from the order acknowledgment (pages 1-2):\n"
                "```\n" + order_text[:5000] + "\n```\n\n"
                "Extract ALL door information and return ONLY a JSON object."
            )
        })
        for b64 in eng_images:
            content.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/png", "data": b64}
            })
    else:
        # Text-only fallback (no engineering sheet images)
        content.append({
            "type": "text",
            "text": (
                "Extract door data from this order text.\n"
                "Return ONLY a JSON object.\n\n"
                "```\n" + order_text[:10000] + "\n```"
            )
        })

    content.append({
        "type": "text",
        "text": """
Return this exact JSON (omit fields that are truly absent — never use null):

{
  "customer": "DORI DOORS",
  "order_number": "111673",
  "po_number": "BET TORAH FRONT",
  "doors": [
    {
      "quantity": 4,
      "width": "35-1/2\\"",
      "height": "84-5/8\\"",
      "thickness": "1-3/4\\"",
      "swing": "LH",
      "jamb_width": "...",
      "gauge": "18",
      "core": "Honeycomb",
      "construction": "Lockseam",
      "door_type": "Flush",
      "reinforcements": ["Full Mortise Continuous Hinge", "Reinforce for Panic Bar"],
      "hardware_set": "HW-1",
      "tag_number": "TAG#1234",
      "notes": "..."
    }
  ]
}

Field rules:
- customer: the contractor/customer company name only
- order_number: the numeric job/order ID (e.g. "111673")
- po_number: PO reference without any "PO." or "PO#" prefix (e.g. "BET TORAH FRONT")
- width, height, thickness: fractional-inch strings (e.g. "35-1/2\\"")
- swing: "LH" or "RH" or "LHR" or "RHR". Check the L.H./R.H. column or swing markings.
- jamb_width: frame/jamb opening size if written on the engineering sheet
- tag_number: ONLY if explicitly written on the engineering sheet (e.g. "TAG#2561"). If the TAG column is blank, omit this field entirely.
- reinforcements: list every reinforcement noted (panic bar, door closer, hinges, etc.)
- If all doors in an order are identical, return ONE entry with correct quantity
- If doors differ, return one entry per distinct specification
"""
    })

    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        },
        json={
            "model": "claude-opus-4-5",
            "max_tokens": 2000,
            "system": SYSTEM,
            "messages": [{"role": "user", "content": content}],
        },
        timeout=60,
    )

    if resp.status_code != 200:
        raise RuntimeError(f"Claude API error {resp.status_code}: {resp.text[:300]}")

    raw = resp.json()["content"][0]["text"].strip()
    raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Try to find JSON object in the response
        start = raw.find("{")
        end   = raw.rfind("}") + 1
        if start != -1 and end > start:
            return json.loads(raw[start:end])
        raise RuntimeError(f"Could not parse JSON from Claude response:\n{raw[:400]}")


# ─────────────────────────────────────────────
#  DOCX GENERATION  (matches T1.doc format)
# ─────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        val = kwargs.get(edge, {})
        if val:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"),   val.get("val", "single"))
            tag.set(qn("w:sz"),    str(val.get("sz", 12)))
            tag.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def add_run(para, text, bold=False, size_pt=10, color_hex=None, italic=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = "Arial"
    if color_hex:
        r, g, b = bytes.fromhex(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def build_tag_cell(cell, door, customer, order_number, po_number):
    """Fill a table cell with a door tag matching T1.doc format."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Cell padding
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    "80")
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # ── Line 1: Company name (red background, white text) ──
    p0 = cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p0._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "7A1C1C")
    pPr.append(shd)
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), "30")
    spc.set(qn("w:after"),  "30")
    pPr.append(spc)
    add_run(p0, "ACME & DORF DOOR CORP", bold=True, size_pt=9, color_hex="FFFFFF")

    # ── Line 2: Customer ──
    p1 = cell.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr1 = p1._p.get_or_add_pPr()
    spc1 = OxmlElement("w:spacing")
    spc1.set(qn("w:before"), "40")
    spc1.set(qn("w:after"),  "20")
    pPr1.append(spc1)
    add_run(p1, customer.upper(), bold=True, size_pt=10)

    # ── Line 3: Order # + PO ──
    p2 = cell.add_paragraph()
    pPr2 = p2._p.get_or_add_pPr()
    spc2 = OxmlElement("w:spacing")
    spc2.set(qn("w:before"), "0")
    spc2.set(qn("w:after"),  "0")
    pPr2.append(spc2)
    add_run(p2, f"{order_number} PO#{po_number}", bold=False, size_pt=9)

    # ── Line 4: Size + Swing ──
    parts = [door.get("width", ""), door.get("height", "")]
    if door.get("thickness"):
        parts.append(door["thickness"])
    size_str = " X ".join(p for p in parts if p)
    swing = door.get("swing", "")
    p3 = cell.add_paragraph()
    pPr3 = p3._p.get_or_add_pPr()
    spc3 = OxmlElement("w:spacing")
    spc3.set(qn("w:before"), "40")
    spc3.set(qn("w:after"),  "20")
    pPr3.append(spc3)
    size_line = f" {size_str}"
    if swing:
        size_line += f" {swing}"
    add_run(p3, size_line, bold=True, size_pt=11)

    # ── Line 5: Jamb (if present) ──
    if door.get("jamb_width"):
        pj = cell.add_paragraph()
        pPrj = pj._p.get_or_add_pPr()
        spcj = OxmlElement("w:spacing")
        spcj.set(qn("w:before"), "0")
        spcj.set(qn("w:after"),  "0")
        pPrj.append(spcj)
        add_run(pj, f" Jamb: {door['jamb_width']}", bold=False, size_pt=8, color_hex="555555")

    # ── Spec pills row: gauge, core, type ──
    specs = []
    if door.get("gauge"):     specs.append(door["gauge"] + " ga.")
    if door.get("core"):      specs.append(door["core"])
    if door.get("door_type"): specs.append(door["door_type"])
    if specs:
        ps = cell.add_paragraph()
        pPrs = ps._p.get_or_add_pPr()
        spcs = OxmlElement("w:spacing")
        spcs.set(qn("w:before"), "0")
        spcs.set(qn("w:after"),  "0")
        pPrs.append(spcs)
        add_run(ps, " " + "  |  ".join(specs), bold=False, size_pt=7.5, color_hex="555555")

    # ── Reinforcements ──
    for r in (door.get("reinforcements") or [])[:3]:
        pr = cell.add_paragraph()
        pPrr = pr._p.get_or_add_pPr()
        spcr = OxmlElement("w:spacing")
        spcr.set(qn("w:before"), "0")
        spcr.set(qn("w:after"),  "0")
        pPrr.append(spcr)
        add_run(pr, f" \u2022 {r}", bold=False, size_pt=7.5, color_hex="444444")

    # ── TAG# line ──
    pt = cell.add_paragraph()
    pPrt = pt._p.get_or_add_pPr()
    spct = OxmlElement("w:spacing")
    spct.set(qn("w:before"), "40")
    spct.set(qn("w:after"),  "0")
    pPrt.append(spct)

    tag_num = door.get("tag_number", "")
    if tag_num:
        add_run(pt, tag_num, bold=True, size_pt=9)
    else:
        add_run(pt, "TAG# ___________", bold=False, size_pt=8, color_hex="999999")


def generate_docx(data, copies=1):
    """Generate a DOCX with door tags in 3-column Avery-style layout."""
    customer     = data.get("customer", "")
    order_number = data.get("order_number", "")
    po_number    = data.get("po_number", "")
    doors        = data.get("doors", [])

    # Expand by quantity × copies
    tags = []
    for door in doors:
        qty = door.get("quantity", 1)
        for _ in range(qty * copies):
            tags.append(door)

    doc = Document()

    # Page setup: Letter, narrow margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin    = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin   = Inches(0.4)
    section.right_margin  = Inches(0.4)

    # Default paragraph spacing
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after  = Pt(0)

    # Table: 3 cols, N rows
    COLS      = 3
    TAG_W_IN  = 2.5   # inches per tag cell
    TAG_H_IN  = 1.5   # row height

    num_rows = max(1, -(-len(tags) // COLS))   # ceiling division
    table = doc.add_table(rows=num_rows, cols=COLS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style     = "Table Grid"

    # Set column widths
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(TAG_W_IN)

    # Set row heights
    for row in table.rows:
        tr   = row._tr
        trPr = tr.get_or_add_trPr()
        trH  = OxmlElement("w:trHeight")
        trH.set(qn("w:val"),  str(int(TAG_H_IN * 1440)))  # DXA
        trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)

    # Fill cells
    for idx, door in enumerate(tags):
        row_i = idx // COLS
        col_i = idx % COLS
        cell  = table.cell(row_i, col_i)

        border_style = {"val": "single", "sz": 12, "color": "000000"}
        set_cell_border(cell, top=border_style, left=border_style,
                        bottom=border_style, right=border_style)
        build_tag_cell(cell, door, customer, order_number, po_number)

    # Empty remaining cells get borders too
    total_cells = num_rows * COLS
    for idx in range(len(tags), total_cells):
        row_i = idx // COLS
        col_i = idx % COLS
        cell  = table.cell(row_i, col_i)
        border_style = {"val": "single", "sz": 12, "color": "CCCCCC"}
        set_cell_border(cell, top=border_style, left=border_style,
                        bottom=border_style, right=border_style)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
#  FLASK ROUTES
# ─────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():
    api_key = request.form.get("api_key", "").strip() or ANTHROPIC_API_KEY
    if not api_key:
        return jsonify({"error": "No API key provided. Set ANTHROPIC_API_KEY or enter it in the form."}), 400

    if "pdf" not in request.files:
        return jsonify({"error": "No PDF file uploaded."}), 400

    pdf_file = request.files["pdf"]
    if not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "Please upload a PDF file."}), 400

    copies = int(request.form.get("copies", 1))

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        pdf_file.save(tmp.name)
        tmp_path = tmp.name

    try:
        # Step 1 — extract
        order_text, eng_images, num_pages = extract_pdf(tmp_path)

        # Step 2 — Claude
        data = extract_with_claude(order_text, eng_images, api_key)

        # Step 3 — DOCX
        docx_buf = generate_docx(data, copies=copies)

        filename = f"tags_{data.get('order_number', 'order')}_{data.get('customer', '').replace(' ', '_')}.docx"

        return send_file(
            docx_buf,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        return jsonify({"error": str(e)}), 500

    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
